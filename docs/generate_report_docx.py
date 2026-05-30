#!/usr/bin/env python3
"""Generate PROJECT_REPORT.docx from PROJECT_REPORT.md (Microsoft Word)."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

DOCS = Path(__file__).resolve().parent
MD_PATH = DOCS / "PROJECT_REPORT.md"
DOCX_PATH = DOCS / "PROJECT_REPORT.docx"


def strip_md(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    return text.replace("&lt;", "<").replace("&gt;", ">").strip()


def parse_blocks(path: Path) -> list[tuple[str, str]]:
    lines = path.read_text(encoding="utf-8").splitlines()
    blocks: list[tuple[str, str]] = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.strip() == "---":
            i += 1
            continue
        if line.startswith("# ") and not line.startswith("## "):
            blocks.append(("title", strip_md(line[2:])))
            i += 1
            continue
        if line.startswith("## "):
            blocks.append(("h1", strip_md(line[3:])))
            i += 1
            continue
        if line.startswith("### "):
            blocks.append(("h2", strip_md(line[4:])))
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                if not re.match(r"^\|[\s\-:|]+\|$", lines[i].strip()):
                    table_lines.append(lines[i])
                i += 1
            blocks.append(("table", "\n".join(table_lines)))
            continue
        if line.startswith("- ") or line.startswith("* "):
            bullets = []
            while i < len(lines) and (
                lines[i].startswith("- ") or lines[i].startswith("* ")
            ):
                bullets.append(strip_md(lines[i][2:]))
                i += 1
            blocks.append(("bullet", "\n".join(bullets)))
            continue
        if line.strip().startswith("```"):
            i += 1
            code = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            blocks.append(("code", "\n".join(code)))
            continue
        if line.strip():
            para = [line]
            i += 1
            while i < len(lines) and lines[i].strip() and not lines[i].startswith("#"):
                if lines[i].startswith("|") or lines[i].startswith("- "):
                    break
                if lines[i].strip() == "---":
                    break
                para.append(lines[i])
                i += 1
            blocks.append(("body", strip_md(" ".join(para))))
            continue
        i += 1
    return blocks


def add_table(doc: Document, raw: str) -> None:
    rows = []
    for line in raw.splitlines():
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(col_count):
            cell = row[ci] if ci < len(row) else ""
            table.rows[ri].cells[ci].text = cell
            if ri == 0:
                for p in table.rows[ri].cells[ci].paragraphs:
                    for run in p.runs:
                        run.bold = True


def build_docx(blocks: list[tuple[str, str]]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    titles = [c for k, c in blocks if k == "title"]
    main = titles[0] if titles else "Cholo BD"
    sub = titles[1] if len(titles) > 1 else "Smart Travel BD"

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(main)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(22, 101, 52)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(sub)
    run.font.size = Pt(14)

    meta_keys = ("Project Type:", "Platform:", "Technology", "Version:", "Academic Year:")
    for k, c in blocks:
        if k == "body" and any(c.startswith(m) or m in c for m in meta_keys):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(c.replace("**", ""))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("University Project Report")
    run.italic = True
    doc.add_page_break()

    skip_meta = True
    for kind, content in blocks:
        if kind == "title":
            continue
        if skip_meta:
            if kind == "h1" and "abstract" in content.lower():
                skip_meta = False
            else:
                continue

        if kind == "h1":
            if content.lower().startswith("table of contents"):
                doc.add_page_break()
            doc.add_heading(content, level=1)
        elif kind == "h2":
            doc.add_heading(content, level=2)
        elif kind == "table":
            add_table(doc, content)
            doc.add_paragraph()
        elif kind == "bullet":
            for item in content.split("\n"):
                doc.add_paragraph(item, style="List Bullet")
        elif kind == "code":
            for line in content.split("\n"):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
            doc.add_paragraph()
        elif kind == "body":
            doc.add_paragraph(content)

    doc.save(str(DOCX_PATH))


def main() -> None:
    if not MD_PATH.exists():
        print(f"Missing {MD_PATH}")
        sys.exit(1)
    blocks = parse_blocks(MD_PATH)
    build_docx(blocks)
    print(f"Generated: {DOCX_PATH}")
    print(f"Size: {DOCX_PATH.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
